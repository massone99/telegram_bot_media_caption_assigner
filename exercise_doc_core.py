#!/usr/bin/env python3
import json
import math
import re
import subprocess
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Iterable


VIDEO_EXTENSIONS = (
    ".mp4",
    ".mkv",
    ".mov",
    ".m4v",
    ".webm",
    ".avi",
    ".mpg",
    ".mpeg",
)

CUE_KEYWORDS = (
    "start",
    "setup",
    "position",
    "stance",
    "place",
    "put",
    "bring",
    "move",
    "hold",
    "reach",
    "extend",
    "flex",
    "rotate",
    "drive",
    "pull",
    "push",
    "squeeze",
    "contract",
    "stretch",
    "breathe",
    "repeat",
    "rep",
    "sets",
    "exercise",
)


@dataclass(frozen=True)
class TranscriptSegment:
    start: float
    end: float
    text: str


@dataclass(frozen=True)
class ExerciseBlock:
    title: str
    start: float
    end: float
    segments: tuple[TranscriptSegment, ...]


@dataclass(frozen=True)
class ScreenshotCue:
    time: float
    reason: str
    score: int
    text: str


@dataclass(frozen=True)
class ScreenshotGroup:
    cue: ScreenshotCue
    segments: tuple[TranscriptSegment, ...]


@dataclass(frozen=True)
class TranscriptDocument:
    transcript_path: Path
    media_path: Path | None
    title: str
    duration: float | None
    blocks: tuple[ExerciseBlock, ...]


def discover_transcripts(path: Path) -> list[Path]:
    if path.is_file():
        if path.suffix.lower() != ".json":
            raise ValueError(f"Transcript must be a .json file: {path}")
        return [path]
    if not path.is_dir():
        raise ValueError(f"Input path not found: {path}")
    return sorted(path.rglob("*.json"))


def load_transcript_payload(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def parse_segments(payload: dict) -> list[TranscriptSegment]:
    segments = []
    for raw in payload.get("segments", []):
        text = str(raw.get("text", "")).strip()
        if not text:
            continue
        segments.append(
            TranscriptSegment(
                start=max(0.0, float(raw.get("start", 0.0))),
                end=max(0.0, float(raw.get("end", raw.get("start", 0.0)))),
                text=text,
            )
        )
    return segments


def infer_media_path(payload: dict, transcript_path: Path, media_root: Path | None = None) -> Path | None:
    source = payload.get("source")
    if source:
        source_path = Path(source)
        if source_path.exists():
            return source_path

    roots = [transcript_path.parent]
    if media_root is not None:
        roots.insert(0, media_root)

    for root in roots:
        for extension in VIDEO_EXTENSIONS:
            candidate = root / f"{transcript_path.stem}{extension}"
            if candidate.exists():
                return candidate
    return None


def title_from_path(path: Path) -> str:
    title = re.sub(r"[_-]+", " ", path.stem)
    title = re.sub(r"\s+", " ", title).strip()
    return title or path.stem


def build_blocks(
    segments: Iterable[TranscriptSegment],
    title: str,
    split_on_pauses: bool = False,
    pause_seconds: float = 12.0,
    min_block_seconds: float = 20.0,
) -> tuple[ExerciseBlock, ...]:
    ordered = sorted(segments, key=lambda segment: (segment.start, segment.end))
    if not ordered:
        return ()

    if not split_on_pauses:
        return (
            ExerciseBlock(
                title=title,
                start=ordered[0].start,
                end=ordered[-1].end,
                segments=tuple(ordered),
            ),
        )

    blocks: list[ExerciseBlock] = []
    current: list[TranscriptSegment] = [ordered[0]]
    for segment in ordered[1:]:
        gap = segment.start - current[-1].end
        current_duration = current[-1].end - current[0].start
        if gap >= pause_seconds and current_duration >= min_block_seconds:
            blocks.append(_block_from_segments(title, len(blocks) + 1, current))
            current = [segment]
        else:
            current.append(segment)
    if current:
        if blocks and current[-1].end - current[0].start < min_block_seconds:
            merged = list(blocks[-1].segments) + current
            blocks[-1] = _block_from_segments(title, len(blocks), merged)
        else:
            blocks.append(_block_from_segments(title, len(blocks) + 1, current))
    return tuple(blocks)


def _block_from_segments(title: str, index: int, segments: list[TranscriptSegment]) -> ExerciseBlock:
    return ExerciseBlock(
        title=f"{title} - Part {index}",
        start=segments[0].start,
        end=segments[-1].end,
        segments=tuple(segments),
    )


def cue_score(text: str) -> int:
    lowered = text.lower()
    score = 0
    for keyword in CUE_KEYWORDS:
        if re.search(rf"\b{re.escape(keyword)}\b", lowered):
            score += 2
    if "?" in text:
        score -= 1
    word_count = len(lowered.split())
    if 4 <= word_count <= 28:
        score += 1
    return score


def screenshot_cues(
    block: ExerciseBlock,
    count: int = 0,
    seconds_per_screenshot: int = 45,
    min_count: int = 1,
    max_count: int = 12,
    edge_padding_seconds: float = 1.0,
) -> tuple[ScreenshotCue, ...]:
    if count < 1:
        count = screenshot_count_for_block(
            block,
            seconds_per_screenshot=seconds_per_screenshot,
            min_count=min_count,
            max_count=max_count,
        )
    if count < 1:
        return ()

    candidates: list[ScreenshotCue] = []
    for segment in block.segments:
        score = cue_score(segment.text)
        if score <= 0:
            continue
        time = midpoint(segment.start, segment.end)
        candidates.append(ScreenshotCue(time=time, reason="transcript-cue", score=score, text=segment.text))

    fallback_times = evenly_spaced_times(block.start, block.end, count, edge_padding_seconds)
    for time in fallback_times:
        candidates.append(ScreenshotCue(time=time, reason="fallback", score=0, text=""))

    selected: list[ScreenshotCue] = []
    for cue in sorted(candidates, key=lambda item: (-item.score, item.time)):
        bounded = ScreenshotCue(
            time=clamp(cue.time, block.start + edge_padding_seconds, block.end - edge_padding_seconds),
            reason=cue.reason,
            score=cue.score,
            text=cue.text,
        )
        if all(abs(bounded.time - existing.time) >= 3.0 for existing in selected):
            selected.append(bounded)
        if len(selected) >= count:
            break

    return tuple(sorted(selected, key=lambda item: item.time))


def screenshot_count_for_block(
    block: ExerciseBlock,
    seconds_per_screenshot: int = 45,
    min_count: int = 1,
    max_count: int = 12,
) -> int:
    duration = max(0.0, block.end - block.start)
    if seconds_per_screenshot < 1:
        seconds_per_screenshot = 45
    count = math.ceil(duration / seconds_per_screenshot)
    if max_count > 0:
        count = min(max_count, count)
    return max(min_count, count)


def group_segments_by_cue(block: ExerciseBlock, cues: Iterable[ScreenshotCue]) -> tuple[ScreenshotGroup, ...]:
    ordered_cues = tuple(sorted(cues, key=lambda cue: cue.time))
    if not ordered_cues:
        return ()
    groups = []
    for index, cue in enumerate(ordered_cues):
        low = block.start if index == 0 else midpoint(ordered_cues[index - 1].time, cue.time)
        high = block.end if index == len(ordered_cues) - 1 else midpoint(cue.time, ordered_cues[index + 1].time)
        segments = tuple(
            segment
            for segment in block.segments
            if segment_intersects_window(segment, low, high)
        )
        groups.append(ScreenshotGroup(cue=cue, segments=segments))
    return tuple(groups)


def segment_intersects_window(segment: TranscriptSegment, start: float, end: float) -> bool:
    segment_midpoint = midpoint(segment.start, segment.end)
    return start <= segment_midpoint <= end


def midpoint(start: float, end: float) -> float:
    return start + max(0.0, end - start) / 2.0


def evenly_spaced_times(start: float, end: float, count: int, edge_padding_seconds: float) -> list[float]:
    if count <= 0:
        return []
    low = start + edge_padding_seconds
    high = end - edge_padding_seconds
    if high <= low:
        return [max(0.0, midpoint(start, end))]
    if count == 1:
        return [midpoint(low, high)]
    step = (high - low) / (count - 1)
    return [low + step * index for index in range(count)]


def clamp(value: float, low: float, high: float) -> float:
    if high < low:
        return max(0.0, (low + high) / 2.0)
    return max(low, min(high, value))


def seconds_to_timestamp(seconds: float) -> str:
    milliseconds = round(max(0.0, seconds) * 1000)
    hours, remainder = divmod(milliseconds, 3_600_000)
    minutes, remainder = divmod(remainder, 60_000)
    secs, milliseconds = divmod(remainder, 1000)
    return f"{hours:02d}:{minutes:02d}:{secs:02d}.{milliseconds:03d}"


def safe_slug(value: str) -> str:
    slug = re.sub(r"[^A-Za-z0-9._-]+", "_", value.strip())
    slug = re.sub(r"_+", "_", slug).strip("._-")
    return slug or "exercise"


def screenshot_filename(block_index: int, cue_index: int, cue: ScreenshotCue) -> str:
    timestamp = seconds_to_timestamp(cue.time).replace(":", "-").replace(".", "-")
    return f"part_{block_index:02d}_cue_{cue_index:02d}_{timestamp}.jpg"


def extract_screenshot(
    media_path: Path,
    output_path: Path,
    timestamp: float,
    ffmpeg_bin: str = "ffmpeg",
    overwrite: bool = False,
) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    if output_path.exists() and not overwrite:
        return
    command = [
        ffmpeg_bin,
        "-y" if overwrite else "-n",
        "-hide_banner",
        "-loglevel",
        "error",
        "-ss",
        f"{timestamp:.3f}",
        "-i",
        str(media_path),
        "-frames:v",
        "1",
        "-q:v",
        "2",
        str(output_path),
    ]
    subprocess.run(command, check=True)


def build_document_from_transcript(
    transcript_path: Path,
    media_root: Path | None = None,
    split_on_pauses: bool = False,
    pause_seconds: float = 12.0,
) -> TranscriptDocument:
    payload = load_transcript_payload(transcript_path)
    segments = parse_segments(payload)
    title = title_from_path(transcript_path)
    blocks = build_blocks(
        segments,
        title=title,
        split_on_pauses=split_on_pauses,
        pause_seconds=pause_seconds,
    )
    duration = payload.get("duration")
    parsed_duration = None if duration is None else float(duration)
    return TranscriptDocument(
        transcript_path=transcript_path,
        media_path=infer_media_path(payload, transcript_path, media_root),
        title=title,
        duration=parsed_duration,
        blocks=blocks,
    )


def render_markdown(document: TranscriptDocument, image_paths: dict[tuple[int, int], Path]) -> str:
    lines = [f"# {document.title}", ""]
    if document.media_path:
        lines.extend([f"Source: `{document.media_path}`", ""])
    for block_index, block in enumerate(document.blocks, start=1):
        lines.extend(
            [
                f"## {block.title}",
                "",
                f"Time: {seconds_to_timestamp(block.start)} - {seconds_to_timestamp(block.end)}",
                "",
            ]
        )
        for cue_index in range(1, 100):
            image_path = image_paths.get((block_index, cue_index))
            if image_path is None:
                break
            lines.extend([f"![Cue {cue_index}]({image_path.as_posix()})", ""])
            cue = image_cue_from_path(block, image_path, cue_index, image_paths, block_index)
            segments = text_segments_for_image(block, cue_index, image_paths, block_index, cue)
            text = " ".join(segment.text.strip() for segment in segments if segment.text.strip())
            if text:
                lines.extend([text, ""])
    return "\n".join(lines).rstrip() + "\n"


def image_cue_from_path(
    block: ExerciseBlock,
    image_path: Path,
    cue_index: int,
    image_paths: dict[tuple[int, int], Path],
    block_index: int,
) -> ScreenshotCue:
    cues = screenshot_cues(block, count=len([key for key in image_paths if key[0] == block_index]))
    if 1 <= cue_index <= len(cues):
        return cues[cue_index - 1]
    return ScreenshotCue(time=midpoint(block.start, block.end), reason="fallback", score=0, text=str(image_path))


def text_segments_for_image(
    block: ExerciseBlock,
    cue_index: int,
    image_paths: dict[tuple[int, int], Path],
    block_index: int,
    cue: ScreenshotCue,
) -> tuple[TranscriptSegment, ...]:
    count = len([key for key in image_paths if key[0] == block_index])
    cues = screenshot_cues(block, count=count) if count else (cue,)
    groups = group_segments_by_cue(block, cues)
    if 1 <= cue_index <= len(groups):
        return groups[cue_index - 1].segments
    return block.segments


def write_markdown_document(path: Path, document: TranscriptDocument, image_paths: dict[tuple[int, int], Path]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(render_markdown(document, image_paths), encoding="utf-8")


def write_docx_document(path: Path, document: TranscriptDocument, image_paths: dict[tuple[int, int], Path]) -> None:
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError as exc:
        raise RuntimeError("python-docx not found. Install it: python -m pip install python-docx") from exc

    doc = Document()
    doc.add_heading(document.title, level=1)
    if document.media_path:
        doc.add_paragraph(f"Source: {document.media_path}")

    for block_index, block in enumerate(document.blocks, start=1):
        doc.add_heading(block.title, level=2)
        doc.add_paragraph(f"Time: {seconds_to_timestamp(block.start)} - {seconds_to_timestamp(block.end)}")
        for cue_index in range(1, 100):
            image_path = image_paths.get((block_index, cue_index))
            if image_path is None:
                break
            resolved_image_path = image_path if image_path.exists() else path.parent / image_path
            if resolved_image_path.exists():
                doc.add_picture(str(resolved_image_path), width=Inches(5.8))
            cue = image_cue_from_path(block, image_path, cue_index, image_paths, block_index)
            segments = text_segments_for_image(block, cue_index, image_paths, block_index, cue)
            transcript = " ".join(segment.text.strip() for segment in segments if segment.text.strip())
            if transcript:
                doc.add_paragraph(transcript)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def write_pdf_document(path: Path, document: TranscriptDocument, image_paths: dict[tuple[int, int], Path]) -> None:
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer
    except ImportError as exc:
        raise RuntimeError("reportlab not found. Install it: python -m pip install reportlab") from exc

    styles = getSampleStyleSheet()
    elements = [Paragraph(document.title, styles["Title"]), Spacer(1, 12)]
    if document.media_path:
        elements.extend([Paragraph(f"Source: {document.media_path}", styles["BodyText"]), Spacer(1, 8)])

    for block_index, block in enumerate(document.blocks, start=1):
        elements.extend(
            [
                Paragraph(block.title, styles["Heading2"]),
                Paragraph(
                    f"Time: {seconds_to_timestamp(block.start)} - {seconds_to_timestamp(block.end)}",
                    styles["BodyText"],
                ),
                Spacer(1, 8),
            ]
        )
        for cue_index in range(1, 100):
            image_path = image_paths.get((block_index, cue_index))
            if image_path is None:
                break
            resolved_image_path = image_path if image_path.exists() else path.parent / image_path
            if resolved_image_path.exists():
                elements.extend([Image(str(resolved_image_path), width=430, height=242), Spacer(1, 8)])
            cue = image_cue_from_path(block, image_path, cue_index, image_paths, block_index)
            segments = text_segments_for_image(block, cue_index, image_paths, block_index, cue)
            transcript = " ".join(segment.text.strip() for segment in segments if segment.text.strip())
            if transcript:
                elements.extend([Paragraph(transcript, styles["BodyText"]), Spacer(1, 12)])

    path.parent.mkdir(parents=True, exist_ok=True)
    SimpleDocTemplate(str(path), pagesize=letter).build(elements)


def manifest_for_document(
    document: TranscriptDocument,
    cues_by_block: dict[int, tuple[ScreenshotCue, ...]],
    image_paths: dict[tuple[int, int], Path],
) -> dict:
    return {
        "transcript": str(document.transcript_path),
        "media": str(document.media_path) if document.media_path else None,
        "title": document.title,
        "duration": document.duration,
        "blocks": [
            {
                "title": block.title,
                "start": block.start,
                "end": block.end,
                "cues": [
                    {
                        **asdict(cue),
                        "image": str(image_paths.get((block_index, cue_index))),
                        "text": " ".join(
                            segment.text.strip()
                            for segment in group_segments_by_cue(
                                block,
                                cues_by_block.get(block_index, ()),
                            )[cue_index - 1].segments
                            if segment.text.strip()
                        )
                        if cue_index <= len(group_segments_by_cue(block, cues_by_block.get(block_index, ())))
                        else cue.text,
                    }
                    for cue_index, cue in enumerate(cues_by_block.get(block_index, ()), start=1)
                ],
            }
            for block_index, block in enumerate(document.blocks, start=1)
        ],
    }


def write_manifest(path: Path, manifest: dict) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(manifest, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")


def is_finite_positive(value: float) -> bool:
    return math.isfinite(value) and value > 0
